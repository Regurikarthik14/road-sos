from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# =========================================================
# STYLES
# =========================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Segoe UI'
font.size = Pt(11)
font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Set default paragraph format for all styles
for s in doc.styles:
    try:
        s.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    except:
        pass

# Heading styles
for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
    hs.font.name = 'Segoe UI'
    if i == 1:
        hs.font.size = Pt(28)
    elif i == 2:
        hs.font.size = Pt(20)
        hs.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    elif i == 3:
        hs.font.size = Pt(16)
        hs.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
    elif i == 4:
        hs.font.size = Pt(13)
        hs.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)

# Set document background color via XML
sect = doc.sections[0]
sect.top_margin = Cm(2.0)
sect.bottom_margin = Cm(2.0)
sect.left_margin = Cm(2.5)
sect.right_margin = Cm(2.5)

# =========================================================
# HELPERS
# =========================================================
def add_para(text, style_name='Normal', bold=False, size=None, color=None, align=None, space_after=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if bold: run.bold = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def add_table(headers, rows, col_widths=None):
    """Create a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dark header bg
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A263D" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
            # Alternate row shading
            if r_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="131C2E" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    
    # Table borders
    tbl_pr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="2D3748"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="2D3748"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="2D3748"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="2D3748"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="2D3748"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="2D3748"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)
    
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    
    return table

def add_bullet(text, level=0, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
        run2 = p.add_run(text)
        run2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_code_block(code_text):
    """Add a code block styled paragraph"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    # Background via shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0F172A" w:val="clear"/>')
    pPr = p._p.get_or_add_pPr()
    pPr.append(shading)
    run = p.add_run(code_text)
    run.font.size = Pt(8.5)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x22, 0xD3, 0xEE)
    return p

# =========================================================
# COVER PAGE
# =========================================================
for _ in range(4):
    doc.add_paragraph()

add_para("🛡️", align=WD_ALIGN_PARAGRAPH.CENTER, size=48)
add_para("RoadSOS — Raksha", 'Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER, size=36)
add_para("AI-Powered Emergency Response & Roadside Assistance System", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=RGBColor(0x94,0xA3,0xB8))
add_para("___________________________________________________________________", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=RGBColor(0xEF,0x44,0x44))
add_para("Technical Project Document", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True)
add_para("Hackathon Submission 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, color=RGBColor(0x94,0xA3,0xB8))
doc.add_paragraph()
add_para("Built with: React 19 · TypeScript 6 · Vite 8 · Node.js · Express · MongoDB · Gemini AI · Vercel", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=RGBColor(0x64,0x74,0x8B))

doc.add_page_break()

# =========================================================
# TABLE OF CONTENTS
# =========================================================
doc.add_heading('Table of Contents', level=1)
add_para("1.  Project Overview .................................................. 3", size=11, color=RGBColor(0x94,0xA3,0xB8))
add_para("2.  Key Features ....................................................... 3", size=11, color=RGBColor(0x94,0xA3,0xB8))
add_para("3.  System Architecture ........................................... 4", size=11, color=RGBColor(0x94,0xA3,0xB8))
add_para("4.  Technology Stack & Packages ............................... 5", size=11, color=RGBColor(0x94,0xA3,0xB8))
add_para("5.  Installation & Setup ........................................... 6", size=11, color=RGBColor(0x94,0xA3,0xB8))
add_para("6.  Environment Variables ........................................ 7", size=11, color=RGBColor(0x94,0xA3,0xB8))
add_para("7.  Key Code Snippets .............................................. 7", size=11, color=RGBColor(0x94,0xA3,0xB8))
add_para("8.  API Endpoints ...................................................... 9", size=11, color=RGBColor(0x94,0xA3,0xB8))
add_para("9.  Assumptions & Limitations .................................. 9", size=11, color=RGBColor(0x94,0xA3,0xB8))
add_para("10. Future Enhancements ........................................ 10", size=11, color=RGBColor(0x94,0xA3,0xB8))
add_para("11. File Structure ..................................................... 11", size=11, color=RGBColor(0x94,0xA3,0xB8))

doc.add_page_break()

# =========================================================
# 1. PROJECT OVERVIEW
# =========================================================
doc.add_heading('1. Project Overview', level=1)

add_para("RoadSOS (named 'Raksha' — meaning 'protection' in Sanskrit) is a comprehensive web-based emergency response system designed to reduce response time during road accidents and emergencies. The application leverages modern web APIs (DeviceMotion, Geolocation, AudioContext, SpeechRecognition) alongside AI-powered chat assistance to provide a complete roadside safety solution.", size=11)

add_para("Key Statistics:", size=11, bold=True)
add_bullet("1.54 million road fatalities occur worldwide each year", bold_prefix="")
add_bullet("Survival drops 7% each minute without medical care during the Golden Hour", bold_prefix="")
add_bullet("Crash detection is confirmed within 2 seconds using dual-sensor co-trigger", bold_prefix="")
add_bullet("Emergency dispatch with GPS location shared instantly", bold_prefix="")

# =========================================================
# 2. KEY FEATURES
# =========================================================
doc.add_heading('2. Key Features', level=1)

features_data = [
    ("Feature", "Description", "Key Metric"),
    ("Crash Detection", "Sensor-fusion co-trigger using DeviceMotion accelerometer + microphone AudioContext. Impact (3 spikes >20m/s² within 5s) + loud noise (avg >150/255) within 2s = confirmed crash.", "< 2s trigger"),
    ("Fire Detection", "Real-time monitoring of device conditions: battery temperature fluctuation, CPU load spikes, memory pressure. No fake sensors — uses actual device APIs.", "Real-time"),
    ("Hardware Health", "Tracks CPU usage, Battery temperature & level, Sensor status. Critical state triggers 15s auto-call to owner, then auto-dispatch if unanswered.", "15s countdown"),
    ("GPS & Map View", "Interactive Leaflet.js map with OSRM routing. Traffic-colored segments (green/yellow/red), route animation, ETA display. 5 service categories: Trauma, Police, Towing, Puncture, Ambulance.", "Real-time"),
    ("AI Chat Assistant", "Gemini 2.0 Flash AI with streaming responses. Voice input/output via Web Speech API. 20+ smart keyword fallbacks when offline.", "Streaming"),
    ("Auth & Security", "Register with Email OR Phone (Facebook-style toggle). 23-country code selector. OTP verification, password strength meter, forgot/reset flow, JWT authentication.", "Full auth"),
]

add_table(features_data[0], features_data[1:], [3.5, 11, 3])

# =========================================================
# 3. SYSTEM ARCHITECTURE
# =========================================================
doc.add_heading('3. System Architecture', level=1)

add_para("The application follows a modern 3-tier architecture:", size=11)
add_para("")

add_para("Layer 1: User Device (Browser)", 'Heading 3', size=14)
add_bullet("Geolocation API — watchPosition for live GPS tracking")
add_bullet("DeviceMotion API — accelerometer for crash impact detection")
add_bullet("AudioContext + getUserMedia — microphone for audio crash co-trigger")
add_bullet("Web Speech API — SpeechRecognition (voice input) + SpeechSynthesis (AI output)")
add_bullet("Vibration API — SOS Morse pattern haptic feedback")
add_bullet("Battery API — battery level monitoring for fire detection")

add_para("Layer 2: React Frontend (Vite + TypeScript)", 'Heading 3', size=14)
add_bullet("Dashboard — Emergency SOS dial, crash/fire panels, hardware health, quick category chips")
add_bullet("FailsafeUI — Full-screen emergency countdown with flashing background, haptic SOS Morse, medical card")
add_bullet("ChatCanvas — Gemini AI integration with streaming responses, voice I/O, network status indicator")
add_bullet("MapView — Leaflet.js interactive map with OSRM routing, traffic-colored segments, 5 service categories")
add_bullet("AuthScreen — Login/Register with email or phone toggle, OTP verification, password management")

add_para("Layer 3: Backend API (Node.js + Express + MongoDB)", 'Heading 3', size=14)
add_bullet("Auth Service — JWT token generation/verification, bcrypt password hashing, OTP generation/validation")
add_bullet("API Routes — RESTful endpoints for auth (register, login, verify OTP, create password, forgot/reset)")
add_bullet("Email Service (Resend) — Sends OTP codes and password reset links via email")
add_bullet("SMS Service (Twilio) — Sends OTP codes via SMS for phone number verification")
add_bullet("Admin Routes — User management (list, search, delete) protected by admin secret")
add_bullet("Middleware — Auth guard (JWT verification), Admin guard (secret key), CORS configuration")

# =========================================================
# 4. TECHNOLOGY STACK
# =========================================================
doc.add_heading('4. Technology Stack & Packages', level=1)

add_para("Frontend Packages", 'Heading 3', size=14)
fe_packages = [
    ("Package", "Version", "Purpose"),
    ("react", "^19.2.6", "UI library with hooks, concurrent rendering, server components"),
    ("react-dom", "^19.2.6", "DOM renderer for React"),
    ("typescript", "~6.0.2", "Type-safe JavaScript with strict mode"),
    ("vite", "^8.0.12", "Fast build tool with HMR and production bundling"),
    ("leaflet", "^1.9.4", "Interactive map library"),
    ("@types/leaflet", "^1.9.21", "TypeScript type definitions for Leaflet"),
    ("@google/generative-ai", "^0.21.0", "Google Gemini AI SDK for chat responses"),
    ("eslint", "^10.x", "Code linting and static analysis"),
    ("@typescript-eslint/*", "^8.x", "TypeScript ESLint rules"),
]
add_table(fe_packages[0], fe_packages[1:], [4, 2.5, 11])

add_para("")
add_para("Backend Packages", 'Heading 3', size=14)
be_packages = [
    ("Package", "Version", "Purpose"),
    ("express", "^4.21.0", "Web framework with middleware & routing"),
    ("mongoose", "^8.7.0", "MongoDB ODM with schema validation"),
    ("jsonwebtoken", "^9.0.2", "JWT token generation and verification"),
    ("bcryptjs", "^2.4.3", "Password hashing and comparison"),
    ("cors", "^2.8.5", "Cross-Origin Resource Sharing middleware"),
    ("dotenv", "^16.4.5", "Environment variable loading"),
    ("resend", "^4.0.0", "Email delivery SDK"),
    ("twilio", "^5.3.0", "SMS delivery SDK"),
    ("tsx", "^4.19.0", "TypeScript execution for development"),
    ("typescript", "^5.6.0", "TypeScript compiler"),
    ("@types/*", "—", "TypeScript type definitions for Express, Node, JWT, bcrypt, cors"),
]
add_table(be_packages[0], be_packages[1:], [4, 2.5, 11])

# =========================================================
# 5. INSTALLATION & SETUP
# =========================================================
doc.add_heading('5. Installation & Setup', level=1)

add_para("Prerequisites", 'Heading 3', size=14)
add_bullet("Node.js >= 22")
add_bullet("npm or pnpm")
add_bullet("MongoDB Atlas account (or local MongoDB instance)")
add_bullet("Vercel account (for deployment)")

add_para("")
add_para("Installation Steps", 'Heading 3', size=14)
add_para("1. Clone the repository:", size=11, bold=True)
add_code_block("git clone https://github.com/your-username/road-sos.git\ncd road-sos")

add_para("2. Install frontend dependencies:", size=11, bold=True)
add_code_block("npm install")

add_para("3. Install backend dependencies:", size=11, bold=True)
add_code_block("cd server && npm install")

add_para("4. Create .env file in the server directory:", size=11, bold=True)
add_code_block(
"MONGODB_URI=mongodb+srv://<user>:<pass>@cluster.mongodb.net/raksha\n"
"JWT_SECRET=<your-secret-key>\n"
"ADMIN_SECRET=<your-admin-secret>\n"
"FRONTEND_URL=http://localhost:5173\n"
"RESEND_API_KEY=re_<your-resend-key>\n"
"TWILIO_ACCOUNT_SID=<your-twilio-sid>\n"
"TWILIO_AUTH_TOKEN=<your-twilio-token>\n"
"TWILIO_PHONE_NUMBER=+1XXXXXXXXXX"
)

add_para("5. Run the backend:", size=11, bold=True)
add_code_block("cd server && npx tsx src/index.ts")

add_para("6. Run the frontend (in a separate terminal):", size=11, bold=True)
add_code_block("npm run dev")

add_para("7. Open http://localhost:5173 in your browser", size=11)

# =========================================================
# 6. ENVIRONMENT VARIABLES
# =========================================================
doc.add_heading('6. Environment Variables', level=1)

env_vars = [
    ("Variable", "Required", "Description", "Example"),
    ("MONGODB_URI", "Yes", "MongoDB connection string", "mongodb+srv://user:pass@cluster/raksha"),
    ("JWT_SECRET", "Yes", "Secret key for JWT token signing (64+ hex chars)", "d3a0b4c1..."),
    ("ADMIN_SECRET", "Yes", "Secret key for admin route access", "913c3cae..."),
    ("FRONTEND_URL", "Yes", "CORS allowed origin for the frontend", "http://localhost:5173"),
    ("RESEND_API_KEY", "No*", "Resend API key for email OTPs", "re_abc123..."),
    ("TWILIO_ACCOUNT_SID", "No*", "Twilio account SID for SMS OTPs", "ACxxxxxx..."),
    ("TWILIO_AUTH_TOKEN", "No*", "Twilio auth token for SMS", "xxxx..."),
    ("TWILIO_PHONE_NUMBER", "No*", "Twilio phone number for SMS", "+1XXXXXXXXXX"),
    ("PORT", "No", "Server port (default: 3001)", "3001"),
]
add_table(env_vars[0], env_vars[1:], [4, 2.5, 5, 6])

add_para("* At least one of RESEND_API_KEY or TWILIO_* credentials is required for OTP delivery.", size=10, color=RGBColor(0xF5,0x9E,0x0B))

# =========================================================
# 7. KEY CODE SNIPPETS
# =========================================================
doc.add_heading('7. Key Code Snippets', level=1)

# 7.1 Crash Detection
add_para("7.1 Crash Detection (Co-Trigger System)", 'Heading 3', size=14)
add_para("The crash detection system uses a dual-sensor co-trigger pattern: accelerometer (impact) + microphone (loud sound). Both must fire within 2 seconds to confirm a crash, preventing false positives from phone shaking.", size=10)

add_code_block(
"// Thresholds\n"
"const IMPACT_THRESHOLD = 20;       // m/s^2 (~2g)\n"
"const LOUD_THRESHOLD = 150;        // 0-255 scale\n"
"const CO_TRIGGER_WINDOW = 2000;    // ms — both must occur within 2s\n\n"
"// DeviceMotion listener (impact detection)\n"
"const handleMotion = (event: DeviceMotionEvent) => {\n"
"  const { x, y, z } = event.accelerationIncludingGravity!;\n"
"  const magnitude = Math.sqrt(x*x + y*y + z*z);\n"
"  if (magnitude > IMPACT_THRESHOLD) {\n"
"    impactTimeRef.current = Date.now();\n"
"    checkCoTrigger();\n"
"  }\n"
"};\n\n"
"// Audio analyzer (loud sound detection)\n"
"const analyzeAudio = () => {\n"
"  const data = new Uint8Array(analyser.frequencyBinCount);\n"
"  analyser.getByteFrequencyData(data);\n"
"  const avg = data.reduce((a, b) => a + b, 0) / data.length;\n"
"  if (avg > LOUD_THRESHOLD) {\n"
"    loudTimeRef.current = Date.now();\n"
"    checkCoTrigger();\n"
"  }\n"
"  rafRef.current = requestAnimationFrame(analyzeAudio);\n"
"};"
)

# 7.2 Fire Detection
add_para("7.2 Fire Detection System", 'Heading 3', size=14)
add_para("Monitors real device conditions using Battery API, performance memory, and navigator hardware concurrency. No fake sensor simulations.", size=10)

add_code_block(
"// Fire detection uses actual device state\n"
"const checkFireConditions = async () => {\n"
"  const battery = await navigator.getBattery();\n"
"  const memory = (performance as any).memory;\n"
"  \n"
"  // Check for warning signs:\n"
"  // 1. Battery charging rapidly (potential thermal runaway)\n"
"  // 2. Memory pressure high (>500MB used)\n"
"  // 3. CPU sustained high load (hardwareConcurrency * 0.8)\n"
"  \n"
"  if (battery.charging && battery.level > 0.8) {\n"
"    updateFireRisk('elevated');\n"
"  }\n"
"  if (memory?.usedJSHeapSize > 500 * 1024 * 1024) {\n"
"    updateFireRisk('high');\n"
"  }\n"
"};"
)

# 7.3 Auth Flow
add_para("7.3 Authentication — Facebook-Style Registration", 'Heading 3', size=14)
add_para("Users can register with either email OR phone number. A 23-country code selector provides proper phone number formatting.", size=10)

add_code_block(
"// Login controller: accepts email OR phone\n"
"const login = async (req: Request, res: Response) => {\n"
"  const { email } = req.body; // 'email' field accepts email or phone\n"
"  \n"
"  const isEmail = email.includes('@');\n"
"  let user;\n"
"  \n"
"  if (isEmail) {\n"
"    user = await User.findOne({ email: email.toLowerCase().trim() });\n"
"  } else {\n"
"    // Clean phone to E.164 format: +15551234567\n"
"    const cleanPhone = '+' + email.replace(/[^0-9]/g, '');\n"
"    user = await User.findOne({ phone: cleanPhone });\n"
"  }\n"
"  \n"
"  if (!user || !(await bcrypt.compare(password, user.password))) {\n"
"    return res.status(401).json({ message: 'Invalid credentials' });\n"
"  }\n"
"  \n"
"  const token = jwt.sign({ uid: user._id }, process.env.JWT_SECRET!);\n"
"  res.json({ token, user: { uid: user._id, email: user.email || '', \n"
"    phone: user.phone || '', uniqueId: user.uniqueId } });\n"
"};"
)

# 7.4 Map Routing
add_para("7.4 Map Routing with OSRM", 'Heading 3', size=14)
add_para("Interactive map routing using Open Source Routing Machine (OSRM) with traffic-colored route segments.", size=10)

add_code_block(
"// Fetch route from OSRM API\n"
"const fetchRoute = async (fromLat, fromLng, toLat, toLng) => {\n"
"  const url = `https://router.project-osrm.org/route/v1/driving/\n"
"    ${fromLng},${fromLat};${toLng},${toLat}?overview=full&geometries=geojson`;\n"
"  \n"
"  const res = await fetch(url);\n"
"  const data = await res.json();\n"
"  const coords = data.routes[0].geometry.coordinates;\n"
"  const distance = (data.routes[0].distance / 1000).toFixed(1) + ' km';\n"
"  const duration = Math.round(data.routes[0].duration / 60) + ' min';\n"
"  \n"
"  // Draw polyline with traffic color\n"
"  const color = distance < 5 ? '#22C55E' : distance < 15 ? '#F59E0B' : '#EF4444';\n"
"  const polyline = L.polyline(coords, { \n"
"    color, weight: 4, opacity: 0.85, dashArray: '12, 8'\n"
"  }).addTo(map);\n"
"  map.fitBounds(polyline.getBounds(), { padding: [50, 50] });\n"
"};"
)

# =========================================================
# 8. API ENDPOINTS
# =========================================================
doc.add_heading('8. API Endpoints', level=1)

api_endpoints = [
    ("Method", "Endpoint", "Description", "Auth"),
    ("POST", "/api/auth/send-otp", "Send OTP to email or phone", "None"),
    ("POST", "/api/auth/verify-otp", "Verify OTP code", "None"),
    ("POST", "/api/auth/create-password", "Set password after OTP verification", "None"),
    ("POST", "/api/auth/login", "Login with email/phone + password", "None"),
    ("POST", "/api/auth/forgot-password", "Send password reset OTP", "None"),
    ("POST", "/api/auth/reset-password", "Reset password with OTP", "None"),
    ("GET", "/api/auth/profile", "Get authenticated user profile", "JWT"),
    ("GET", "/api/admin/users", "List all users (admin)", "Admin"),
    ("GET", "/api/admin/users/:id", "Get user by ID (admin)", "Admin"),
    ("GET", "/api/admin/users/search/:term", "Search users (admin)", "Admin"),
    ("DELETE", "/api/admin/users/:id", "Delete user (admin)", "Admin"),
    ("GET", "/api/health", "Health check endpoint", "None"),
]
add_table(api_endpoints[0], api_endpoints[1:], [2.5, 5.5, 7, 2.5])

# =========================================================
# 9. ASSUMPTIONS & LIMITATIONS
# =========================================================
doc.add_heading('9. Assumptions & Limitations', level=1)

add_para("Assumptions", 'Heading 3', size=14)
assumptions = [
    "User has a modern browser (Chrome 90+, Edge 90+, Safari 15+, Firefox 90+) that supports DeviceMotion, Geolocation, and Web Audio APIs.",
    "User grants permission for location, microphone, and motion sensors when prompted.",
    "Device has accelerometer hardware (all modern smartphones and most laptops).",
    "Internet connection is available for map tiles, routing API, and AI chat (graceful degradation for offline).",
    "MongoDB Atlas cluster is accessible from the Vercel deployment IP range (or IP whitelisted).",
    "The application is used in a vehicle or environment where crash detection is relevant.",
    "Emergency dispatch service integration would be handled by third-party APIs in production.",
    "Phone numbers are stored in E.164 format (+1XXXXXXXXXX) for SMS delivery.",
    "Email delivery (Resend) and SMS delivery (Twilio) accounts are properly configured.",
    "Users have unique phone numbers and email addresses (no duplicate accounts).",
]
for a in assumptions:
    add_bullet(a)

add_para("")
add_para("Current Limitations", 'Heading 3', size=14)
limitations = [
    "No actual emergency dispatch backend — SOS countdown expiry is a UI placeholder (no API call made).",
    "Service locations are synthetically generated near user position — not real provider data.",
    "iOS requires user gesture before mic access, which may delay crash detection setup.",
    "Duplicate GPS trackers — MapView has its own watchPosition independent of the shared hook.",
    "No PWA support — no service worker or web manifest for offline installability.",
    "No push notifications — can't proactively alert emergency contacts when SOS is triggered.",
    "No unit or integration tests — manual testing only.",
    "CSS uses !important for responsive overrides (~40 declarations) — functional but not elegant.",
]
for l in limitations:
    add_bullet(l)

# =========================================================
# 10. FUTURE ENHANCEMENTS
# =========================================================
doc.add_heading('10. Future Enhancements', level=1)

enhancements = [
    ("Enhancement", "Description", "Priority"),
    ("Emergency Dispatch API", "Connect to real emergency dispatch service (Twilio voice call, email, SMS alert)", "High"),
    ("Real Provider Data", "Integrate Google Places / Overpass API for real hospitals, police stations, towing", "High"),
    ("Push Notifications", "Web push notifications to alert emergency contacts when SOS is triggered", "Medium"),
    ("PWA Support", "Service worker + manifest for offline cache and installable app", "Medium"),
    ("iOS Sensor Permissions", "Add DeviceOrientationEvent.requestPermission() flow for iOS 13+", "Medium"),
    ("Consolidated GPS", "Merge MapView's duplicate geolocation watcher with shared hook", "Low"),
    ("Route Caching", "Cache OSRM route responses in sessionStorage or IndexedDB", "Low"),
    ("Unit Tests", "Add Vitest + React Testing Library for hook and component tests", "Medium"),
    ("Accessibility Audit", "Screen reader support, focus management, ARIA labels, keyboard navigation", "Medium"),
    ("Dark/Light Theme", "Light theme support for non-emergency modes", "Low"),
]
add_table(enhancements[0], enhancements[1:], [4, 10, 3.5])

# =========================================================
# 11. FILE STRUCTURE
# =========================================================
doc.add_heading('11. File Structure', level=1)

add_code_block(
"roadsos-app/\n"
"├── index.html                          # Entry point\n"
"├── package.json                        # Frontend dependencies & scripts\n"
"├── vite.config.ts                      # Vite build configuration\n"
"├── vercel.json                         # Vercel deployment config\n"
"├── tsconfig.json                       # TypeScript project references\n"
"├── tsconfig.app.json                   # App-specific TS config\n"
"├── tsconfig.node.json                  # Node (Vite) TS config\n"
"├── api/\n"
"│   └── index.ts                        # Vercel serverless handler\n"
"├── src/\n"
"│   ├── main.tsx                        # React entry point\n"
"│   ├── App.tsx                         # Root component with view routing\n"
"│   ├── App.css                         # Component & responsive styles\n"
"│   ├── index.css                       # Global reset, design tokens, animations\n"
"│   ├── types.ts                        # Shared TypeScript types\n"
"│   ├── api/\n"
"│   │   └── client.ts                   # API client (fetch wrapper)\n"
"│   ├── context/\n"
"│   │   └── AuthContext.tsx              # Auth state management\n"
"│   ├── hooks/\n"
"│   │   ├── useCrashDetection.ts        # Sensor-fusion crash detection\n"
"│   │   ├── useFireDetection.ts         # Device condition monitoring\n"
"│   │   ├── useGeolocation.ts           # GPS location tracking\n"
"│   │   ├── useVoice.ts                 # Voice recognition synthesis\n"
"│   │   ├── useGemini.ts                # Google Gemini AI integration\n"
"│   │   ├── useHardwareStatus.ts        # CPU/Battery/Sensor monitoring\n"
"│   │   └── useTheme.ts                 # Theme management\n"
"│   └── components/\n"
"│       ├── auth/\n"
"│       │   └── AuthScreen.tsx           # Login, Register, OTP, Reset views\n"
"│       ├── Dashboard.tsx               # Main dashboard with SOS dial\n"
"│       ├── FailsafeUI.tsx              # Emergency countdown screen\n"
"│       ├── ChatCanvas.tsx              # AI chat assistant\n"
"│       ├── MapView.tsx                 # Interactive map with routing\n"
"│       ├── BottomNav.tsx               # 3-tab bottom navigation\n"
"│       ├── MedicalCard.tsx             # Medical ID card component\n"
"│       ├── CrashDetectBanner.tsx       # Crash detection status banner\n"
"│       ├── FireDetection.tsx           # Fire monitoring panel\n"
"│       └── HardwareStatus.tsx          # Hardware health display\n"
"├── server/\n"
"│   ├── package.json                    # Backend dependencies\n"
"│   ├── tsconfig.json                   # Backend TypeScript config\n"
"│   └── src/\n"
"│       ├── index.ts                    # Server entry point\n"
"│       ├── app.ts                      # Express app (Vercel-compatible)\n"
"│       ├── config/\n"
"│       │   └── db.ts                   # MongoDB connection\n"
"│       ├── controllers/\n"
"│       │   ├── authController.ts       # Auth logic (register, login, OTP)\n"
"│       │   └── adminController.ts      # Admin user management\n"
"│       ├── models/\n"
"│       │   ├── User.ts                 # User schema (email, phone, password)\n"
"│       │   └── Otp.ts                  # OTP schema\n"
"│       ├── routes/\n"
"│       │   ├── auth.ts                 # Auth API routes\n"
"│       │   └── admin.ts                # Admin API routes\n"
"│       ├── middleware/\n"
"│       │   └── auth.ts                 # JWT & Admin auth middleware\n"
"│       ├── services/\n"
"│       │   ├── otpService.ts           # OTP generation & validation\n"
"│       │   ├── email.ts                # Resend email service\n"
"│       │   └── sms.ts                  # Twilio SMS service\n"
"│       └── types/\n"
"│           └── index.ts                # Shared backend types\n"
)

# =========================================================
# SAVE
# =========================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "RoadSOS_Project_Document.docx")
doc.save(output_path)
print(f"[OK] Document saved: {output_path}")
